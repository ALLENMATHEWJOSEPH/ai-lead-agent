# exporter.py
import pandas as pd
from docx import Document
import uuid

def export_to_excel(leads):
    df = pd.DataFrame(leads)
    path = f"leads_{uuid.uuid4().hex}.xlsx"
    df.to_excel(path, index=False)
    return path

def export_to_word(leads):
    doc = Document()
    doc.add_heading("Lead Summary Report", 0)

    for idx, lead in enumerate(leads, start=1):
        doc.add_heading(f"Lead {idx}: {lead['Company Name']}", level=1)
        doc.add_paragraph(f"Website: {lead['Website']}")
        doc.add_paragraph(f"CEO: {lead['CEO Name']}")
        doc.add_paragraph(f"Email: {lead['CEO Email']}")
        doc.add_paragraph(f"Phone: {lead['CEO Phone']}")
        doc.add_paragraph(f"CFO: {lead['CFO Name']}")
        doc.add_paragraph(f"Email: {lead['CFO Email']}")
        doc.add_paragraph(f"Phone: {lead['CFO Phone']}")
        doc.add_paragraph(f"Contact Person: {lead['Contact Person Name']}")
        doc.add_paragraph(f"Email: {lead['Contact Email']}")
        doc.add_paragraph(f"Company Summary: {lead['Company Summary']}")
        doc.add_paragraph("---")

    path = f"summary_{uuid.uuid4().hex}.docx"
    doc.save(path)
    return path

